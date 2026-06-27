#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ═══════════════════════════════════════════
# STYLES
# ═══════════════════════════════════════════
sn = doc.styles['Normal']
sn.font.name = 'Times New Roman'
sn.font.size = Pt(12)
sn.paragraph_format.line_spacing = 1.5
sn.paragraph_format.space_after = Pt(0)
sn.paragraph_format.space_before = Pt(0)
sn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for sname, sz, sb, sa in [('Heading 1', 14, 18, 6), ('Heading 2', 12, 12, 6)]:
    s = doc.styles[sname]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = None
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_before = Pt(sb)
    s.paragraph_format.space_after = Pt(sa)
    for rPr in s.element.findall(qn('w:rPr')):
        c = rPr.find(qn('w:color'))
        if c is not None:
            rPr.remove(c)

for section in doc.sections:
    section.top_margin = Cm(4)
    section.right_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)


# ═══════════════════════════════════════════
# NUMBERING: one abstractNum, new numId per restart
# ═══════════════════════════════════════════

_next_abstract_id = [200]
_next_num_id = [200]
_current_num_id = [None]
_current_letter_id = [None]

def _get_numbering_elm():
    return doc.part.numbering_part.element

def _make_abstract_decimal():
    """Create abstractNum for decimal 1. 2. 3. numbering, return abstractNumId."""
    elm = _get_numbering_elm()
    aid = str(_next_abstract_id[0])
    _next_abstract_id[0] += 1

    ab = OxmlElement('w:abstractNum')
    ab.set(qn('w:abstractNumId'), aid)

    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    for tag, val in [('w:start', '1'), ('w:numFmt', 'decimal'), ('w:lvlText', '%1.')]:
        e = OxmlElement(tag)
        e.set(qn('w:val'), val)
        lvl.append(e)
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:firstLine'), '360')
    pPr.append(ind)
    lvl.append(pPr)
    ab.append(lvl)
    elm.append(ab)
    return aid

def _make_abstract_letter():
    """Create abstractNum for lowerLetter a. b. c. numbering, return abstractNumId."""
    elm = _get_numbering_elm()
    aid = str(_next_abstract_id[0])
    _next_abstract_id[0] += 1

    ab = OxmlElement('w:abstractNum')
    ab.set(qn('w:abstractNumId'), aid)

    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    for tag, val in [('w:start', '1'), ('w:numFmt', 'lowerLetter'), ('w:lvlText', '%1.')]:
        e = OxmlElement(tag)
        e.set(qn('w:val'), val)
        lvl.append(e)
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:firstLine'), '360')
    pPr.append(ind)
    lvl.append(pPr)
    ab.append(lvl)
    elm.append(ab)
    return aid

def _make_num(abstract_id, start_at=1):
    """Create a num element referencing abstractNum, with startOverride. Returns numId."""
    elm = _get_numbering_elm()
    nid = str(_next_num_id[0])
    _next_num_id[0] += 1

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), nid)
    ref = OxmlElement('w:abstractNumId')
    ref.set(qn('w:val'), abstract_id)
    num.append(ref)

    override = OxmlElement('w:lvlOverride')
    override.set(qn('w:ilvl'), '0')
    so = OxmlElement('w:startOverride')
    so.set(qn('w:val'), str(start_at))
    override.append(so)
    num.append(override)

    elm.append(num)
    return nid

# Pre-create abstracts
_abstract_decimal = _make_abstract_decimal()
_abstract_letter = _make_abstract_letter()

def numbered(text, restart=False):
    """Numbered item. If restart=True, create new numId starting from 1."""
    if restart:
        _current_num_id[0] = _make_num(_abstract_decimal, 1)
    elif _current_num_id[0] is None:
        _current_num_id[0] = _make_num(_abstract_decimal, 1)
    nid = _current_num_id[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    pPr = p._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)
    nId = OxmlElement('w:numId')
    nId.set(qn('w:val'), nid)
    numPr.append(nId)
    pPr.insert(0, numPr)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def lettered(text):
    """Sub-level lettered item a. b. c."""
    if _current_letter_id[0] is None:
        _current_letter_id[0] = _make_num(_abstract_letter, 1)
    nid = _current_letter_id[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(2.54)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    pPr = p._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)
    nId = OxmlElement('w:numId')
    nId.set(qn('w:val'), nid)
    numPr.append(nId)
    pPr.insert(0, numPr)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


# ═══════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════

def center(text, size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = None

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = None

def para(text, bold=False, italic=False, indent_first=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

def code_block(code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)

def cell_font(cell, text, bold=False, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    p.paragraph_format.line_spacing = 1.15

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell_font(t.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell_font(t.rows[ri + 1].cells[ci], val)


# ═══════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

center('LAPORAN TUGAS AKHIR', 16, True)
center('[ Pemrograman Mobile ]', 12, False, True)
doc.add_paragraph()
center('"Bank Sampah"', 16, True)
for _ in range(3):
    doc.add_paragraph()

center('Disusun oleh:')
center('Muhammad Iqbal Ramadhan')
center('NIM. 231011400285', 12, False, True)
center('Objek/Domain: Pengelolaan Setoran Sampah', 12, False, True)
center('Kelas: Mobile Programming', 12, False, True)

for _ in range(2):
    doc.add_paragraph()

center('PROGRAM STUDI TEKNIK INFORMATIKA', 12, True)
center('FAKULTAS TEKNIK', 12, True)
center('UNIVERSITAS PAMULANG', 12, True)
doc.add_paragraph()
center('2026')
doc.add_page_break()


# ═══════════════════════════════════════════
# KATA PENGANTAR
# ═══════════════════════════════════════════
center('KATA PENGANTAR', 14, True)

para('Segala puji bagi Tuhan Yang Maha Esa yang telah memberikan rahmat, hidayah, dan karunia-Nya sehingga penulis dapat menyelesaikan laporan tugas akhir ini dengan baik. Laporan ini disusun sebagai bagian dari tugas akhir mata kuliah Pemrograman Mobile di Program Studi Teknik Informatika, Universitas Pamulang.')
para('Penulis mengucapkan terima kasih kepada dosen pengampu mata kuliah Pemrograman Mobile yang telah memberikan bimbingan, ilmu, dan motivasi selama perkuliahan berlangsung. Selain itu, penulis juga berterima kasih kepada rekan-rekan mahasiswa yang telah membantu dalam proses pengembangan aplikasi ini.')
para('Penulis menyadari bahwa laporan ini masih memiliki kekurangan, baik dari segi isi maupun penyusunannya. Oleh karena itu, penulis mengharapkan kritik dan saran yang membangun demi perbaikan di masa yang akan datang.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(24)
r = p.add_run('Pamulang, Juni 2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.line_spacing = 1.5
r = p.add_run('Penulis')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
doc.add_page_break()


# ═══════════════════════════════════════════
# DAFTAR ISI
# ═══════════════════════════════════════════
center('DAFTAR ISI', 14, True)
para('Daftar Isi otomatis --- klik kanan di area ini lalu pilih "Update Field" untuk memunculkan/memperbarui.', indent_first=False)
doc.add_page_break()


# ═══════════════════════════════════════════
# BAB I
# ═══════════════════════════════════════════
h1('BAB I PENDAHULUAN')

h2('1.1 Latar Belakang')
para('Bank sampah merupakan salah satu upaya pengelolaan sampah yang melibatkan partisipasi masyarakat dalam memilah dan menyetorkan sampah yang telah terpilah ke suatu tempat penampungan. Konsep bank sampah mulai berkembang di Indonesia sebagai solusi atas permasalahan limbah domestik yang semakin meningkat. Melalui bank sampah, masyarakat dapat menyetorkan sampah seperti plastik, kertas, kaca, logam, dan organik, yang kemudian ditimbang dan dikonversi menjadi nilai saldo atau tabungan.')
para('Pencatatan data setoran sampah secara manual berbasis kertas rentan terhadap kehilangan data, kesalahan pencatatan, dan ketidakefisienan dalam pengelolaan. Oleh karena itu, diperlukan sebuah aplikasi berbasis mobile yang dapat mempermudah proses pencatatan, pengelolaan, dan pencarian data setoran sampah secara digital. Aplikasi ini dikembangkan menggunakan bahasa pemrograman Java dengan framework Android dan penyimpanan data berbasis SQLite.')
para('Laporan tugas akhir ini membahas proses perancangan dan pengembangan aplikasi "Bank Sampah" yang memungkinkan pengguna untuk melakukan operasi CRUD (Create, Read, Update, Delete) pada data setoran sampah, dilengkapi dengan fitur pencarian dan pengurutan data.')

h2('1.2 Rumusan Masalah')
para('Berdasarkan latar belakang di atas, rumusan masalah dalam laporan ini adalah sebagai berikut:')
numbered('Bagaimana merancang dan mengembangkan aplikasi mobile untuk pengelolaan data setoran sampah pada bank sampah?', restart=True)
numbered('Bagaimana mengimplementasikan operasi CRUD pada data setoran sampah menggunakan SQLite sebagai media penyimpanan?')
numbered('Bagaimana efektivitas aplikasi dalam mempermudah pencatatan dan pencarian data setoran sampah?')

h2('1.3 Tujuan')
para('Tujuan dari pembuatan aplikasi ini adalah:')
numbered('Merancang dan mengembangkan aplikasi mobile Bank Sampah yang dapat digunakan untuk mengelola data setoran sampah.', restart=True)
numbered('Mengimplementasikan operasi CRUD (Create, Read, Update, Delete) pada data setoran sampah.')
numbered('Menyediakan fitur pencarian dan pengurutan data untuk mempermudah pengguna dalam menemukan data tertentu.')

h2('1.4 Manfaat')
para('Manfaat yang dapat diperoleh dari aplikasi ini adalah:')
numbered('Bagi Pengguna: Aplikasi ini mempermudah proses pencatatan dan pengelolaan data setoran sampah secara digital, sehingga data lebih terorganisir dan mudah diakses.', restart=True)
numbered('Bagi Penulis: Mengembangkan pemahaman dalam penerapan pemrograman mobile menggunakan Java dan Android, khususnya dalam implementasi CRUD, SQLite, dan Material Design.')

h2('1.5 Batasan & Objek (per-NIM)')
para('Objek dari aplikasi ini adalah pengelolaan data setoran sampah pada bank sampah. Batasan dari aplikasi ini meliputi:')
numbered('Jenis data yang dikelola terdiri dari: nama anggota, jenis sampah (Plastik, Kertas, Botol Kaca, Logam, Organik), berat setoran, dan nilai saldo.', restart=True)
numbered('Penyimpanan data menggunakan SQLite (lokal, tanpa server).')
numbered('Aplikasi hanya dapat diakses melalui perangkat Android dengan minimal SDK 29 (Android 10).')
numbered('Autentikasi pengguna menggunakan mekanisme hardcoded credential (admin/12345).')

doc.add_page_break()


# ═══════════════════════════════════════════
# BAB II
# ═══════════════════════════════════════════
h1('BAB II LANDASAN TEORI')

h2('2.1 Bahasa Pemrograman & IDE')
para('Aplikasi ini dikembangkan menggunakan bahasa pemrograman Java yang merupakan salah satu bahasa pemrograman berorientasi objek (Object-Oriented Programming/OOP) paling populer dan banyak digunakan dalam pengembangan aplikasi Android. Java menawarkan fitur-fitur seperti enkapsulasi, pewarisan (inheritance), dan polimorfisme yang memungkinkan pengembangan aplikasi yang terstruktur dan mudah dirawat.')
para('Integrated Development Environment (IDE) yang digunakan adalah Android Studio, yaitu IDE resmi dari Google untuk pengembangan aplikasi Android. Android Studio menyediakan fitur lengkap seperti code editor, debugger, emulator, dan integrasi dengan Gradle build system yang mempermudah proses pengembangan dan pengujian aplikasi.')

h2('2.2 Konsep yang Diterapkan')
para('Beberapa konsep pemrograman yang diterapkan dalam aplikasi ini antara lain:')
numbered('Pemrograman Berorientasi Objek (OOP): Aplikasi ini menerapkan konsep OOP melalui definisi kelas-kelas seperti Setoran (model data), DatabaseHelper (akses database), LoginActivity, MainActivity, dan AddEditSetoranActivity (komponen antarmuka). Setiap kelas memiliki tanggung jawab yang spesifik sesuai prinsip Single Responsibility.', restart=True)
numbered('Activity Lifecycle: Aplikasi Android berjalan berdasarkan siklus hidup Activity (onCreate, onStart, onResume, onPause, onStop, onDestroy). Aplikasi ini menerapkan lifecycle melalui metode onCreate() pada setiap Activity untuk inisialisasi komponen UI dan data.')
numbered('RecyclerView: Digunakan untuk menampilkan daftar data setoran secara efisien. RecyclerView menggunakan pola ViewHolder untuk mengurangi overhead memori saat scrolling daftar data yang panjang.')
numbered('SharedPreferences: Digunakan untuk menyimpan status login pengguna secara lokal. Status ini berisi data boolean yang menandakan apakah pengguna telah melakukan autentikasi atau belum.')
numbered('Material Design: Antarmuka aplikasi mengikuti pedoman Material Design dari Google yang menggunakan komponen seperti MaterialToolbar, ExtendedFloatingActionButton, TextInputLayout, dan Chip untuk memberikan pengalaman visual yang konsisten.')

h2('2.3 Penyimpanan Data')
para('Aplikasi ini menggunakan SQLite sebagai media penyimpanan data. SQLite adalah sistem manajemen basis data relasional yang tertanam (embedded) dan tidak memerlukan server terpisah. SQLite sangat cocok digunakan pada aplikasi mobile karena ringan, mandiri, dan mendukung query SQL standar.')
para('Database yang digunakan bernama "db_sampah_231011400285.db" dengan satu tabel utama bernama "setoran" yang memiliki kolom: id (INTEGER PRIMARY KEY AUTOINCREMENT), nama_anggota (TEXT), jenis_sampah (TEXT), berat (REAL), dan saldo (REAL). Operasi CRUD dilakukan melalui kelas DatabaseHelper yang menggunakan kelas SQLiteDatabase dan ContentValues dari Android SDK.')

doc.add_page_break()


# ═══════════════════════════════════════════
# BAB III
# ═══════════════════════════════════════════
h1('BAB III ANALISIS & PERANCANGAN')

h2('3.1 Analisis Kebutuhan')
para('Kebutuhan fungsional aplikasi ini meliputi:')
numbered('Autentikasi pengguna: Pengguna dapat melakukan login menggunakan username dan password.', restart=True)
numbered('Menambah data setoran: Pengguna dapat menambahkan catatan setoran sampah baru.')
numbered('Melihat daftar setoran: Pengguna dapat melihat seluruh data setoran yang tersimpan.')
numbered('Mengubah data setoran: Pengguna dapat mengedit data setoran yang sudah ada.')
numbered('Menghapus data setoran: Pengguna dapat menghapus data setoran tertentu.')
numbered('Mencari data: Pengguna dapat mencari data berdasarkan nama anggota.')
numbered('Mengurutkan data: Pengguna dapat mengurutkan data berdasarkan nama, berat, atau saldo.')
numbered('Melihat statistik: Pengguna dapat melihat ringkasan jumlah anggota, total berat, dan total saldo.')

para('Kebutuhan non-fungsional:')

numbered('Antarmuka yang intuitif dan mudah digunakan.', restart=True)
numbered('Data tersimpan secara lokal dan persisten.')
numbered('Aplikasi berjalan pada perangkat Android minimal SDK 29.')

h2('3.2 Rancangan Data')
para('Struktur tabel "setoran" dalam database SQLite adalah sebagai berikut:')

add_table(
    ['Nama Field', 'Tipe Data', 'Keterangan'],
    [
        ['id', 'INTEGER PRIMARY KEY AUTOINCREMENT', 'Identifier unik setiap catatan'],
        ['nama_anggota', 'TEXT NOT NULL', 'Nama anggota yang menyetorkan sampah'],
        ['jenis_sampah', 'TEXT NOT NULL', 'Kategori sampah (Plastik, Kertas, dll)'],
        ['berat', 'REAL NOT NULL', 'Berat setoran dalam kilogram'],
        ['saldo', 'REAL NOT NULL', 'Nilai saldo yang diperoleh'],
    ]
)

para('')
para('Kelas model Setoran memiliki lima atribut: id (int), namaAnggota (String), jenisSampah (String), berat (double), dan saldo (double). Kelas ini menyediakan konstruktor default, konstruktor dengan parameter lengkap, serta metode getter dan setter untuk setiap atribut.')

h2('3.3 Rancangan Antarmuka / Alur Program')
para('Alur program aplikasi dimulai dari LoginActivity sebagai titik masuk (entry point). Pengguna memasukkan username dan password, kemudian sistem memverifikasi kredensial. Jika valid, pengguna dialihkan ke MainActivity. Jika tidak valid, muncul pesan kesalahan.')
para('Pada MainActivity, pengguna dapat melihat daftar setoran dalam bentuk kartu (card view) menggunakan RecyclerView. Tersedia bilah pencarian (search bar) dan spinner pengurutan (sort). Tombol ExtendedFloatingActionButton (FAB) pada pojok kanan bawah digunakan untuk menambah catatan baru. Setiap item memiliki tombol edit dan hapus.')
para('Ketika pengguna menekan tombol tambah atau edit, aplikasi membuka AddEditSetoranActivity yang berisi formulir input dengan field: nama anggota (TextInputEditText), kategori sampah (Spinner), berat setoran (TextInputEditText), dan nilai saldo (TextInputEditText). Validasi input dilakukan sebelum data disimpan ke database.')

doc.add_page_break()


# ═══════════════════════════════════════════
# BAB IV
# ═══════════════════════════════════════════
h1('BAB IV IMPLEMENTASI & PENGUJIAN')

h2('4.1 Lingkungan Implementasi')

add_table(
    ['Komponen', 'Spesifikasi'],
    [
        ['Sistem Operasi', 'CachyOS (Arch Linux) / Windows'],
        ['IDE', 'Android Studio'],
        ['Bahasa Pemrograman', 'Java 11'],
        ['Android SDK', 'API 29 (min) - API 36 (target)'],
        ['Build System', 'Gradle 9.4.1 / AGP 9.2.1'],
        ['Database', 'SQLite (embedded)'],
        ['Library', 'AndroidX AppCompat, Material Design 1.14.0, ConstraintLayout 2.2.1'],
    ]
)

h2('4.2 Implementasi Kode (Bagian Penting)')
para('Berikut adalah penjelasan implementasi kunci dari aplikasi ini:')

para('a. Model Data (Setoran.java)', bold=True)
para('Kelas Setoran merupakan Plain Old Java Object (POJO) yang merepresentasikan satu record data setoran. Kelas ini memiliki lima atribut dan menyediakan konstruktor serta metode getter/setter untuk setiap atribut.')
code_block('''public class Setoran {
    private int id;
    private String namaAnggota;
    private String jenisSampah;
    private double berat;
    private double saldo;

    public Setoran() {}
    public Setoran(int id, String namaAnggota, String jenisSampah,
                   double berat, double saldo) { ... }
    public Setoran(String namaAnggota, String jenisSampah,
                   double berat, double saldo) { ... }

    public int getId() { return id; }
    public String getNamaAnggota() { return namaAnggota; }
    public String getJenisSampah() { return jenisSampah; }
    public double getBerat() { return berat; }
    public double getSaldo() { return saldo; }
}''')

para('b. Akses Database (DatabaseHelper.java)', bold=True)
para('Kelas DatabaseHelper extends SQLiteOpenHelper dan bertanggung jawab atas seluruh operasi CRUD. Kelas ini membuat tabel "setoran" pada metode onCreate(), menyediakan metode insertSetoran(), getAllSetoran(), getSetoranById(), updateSetoran(), deleteSetoran(), dan searchSetoran(). Selain itu, metode insertSeedData() menyisipkan 8 data awal saat database pertama kali dibuat.')
code_block('''public class DatabaseHelper extends SQLiteOpenHelper {
    private static final String DATABASE_NAME = "db_sampah_231011400285.db";
    private static final String TABLE_NAME = "setoran";

    public long insertSetoran(Setoran setoran) {
        SQLiteDatabase db = this.getWritableDatabase();
        ContentValues cv = new ContentValues();
        cv.put(COL_NAMA, setoran.getNamaAnggota());
        cv.put(COL_JENIS, setoran.getJenisSampah());
        cv.put(COL_BERAT, setoran.getBerat());
        cv.put(COL_SALDO, setoran.getSaldo());
        long id = db.insert(TABLE_NAME, null, cv);
        db.close();
        return id;
    }

    public List<Setoran> searchSetoran(String query) {
        Cursor cursor = db.query(TABLE_NAME, null,
            COL_NAMA + " LIKE ?",
            new String[]{"%" + query + "%"}, ...);
        ...
    }
}''')

para('c. Aktivitas Login (LoginActivity.java)', bold=True)
para('LoginActivity merupakan titik masuk aplikasi (launcher activity). Aktivitas ini memeriksa status login dari SharedPreferences. Jika pengguna sudah login, ia langsung dialihkan ke MainActivity. Jika belum, ditampilkan form login dengan dua field: username dan password. Kredensial yang digunakan adalah admin/12345.')

para('d. Aktivitas Utama (MainActivity.java)', bold=True)
para('MainActivity menampilkan daftar setoran menggunakan RecyclerView dengan adapter kustom (SetoranAdapter). Fitur pencarian diimplementasikan menggunakan TextWatcher pada EditText yang memanggil metode searchSetoran() secara real-time. Spinner pengurutan memungkinkan pengguna mengurutkan data berdasarkan nama, berat, atau saldo dalam urutan ascending/descending. Statistik ringkasan (jumlah anggota, total berat, total saldo) dihitung dan ditampilkan secara dinamis.')

para('e. Aktivitas Form Tambah/Edit (AddEditSetoranActivity.java)', bold=True)
para('AddEditSetoranActivity digunakan baik untuk menambah data baru maupun mengedit data yang sudah ada. Aktivitas ini memeriksa intent extra untuk menentukan mode edit atau tambah. Validasi input dilakukan pada semua field: nama tidak boleh kosong, jenis sampah harus dipilih dari spinner, berat dan saldo harus lebih dari nol. Data yang valid disimpan ke database melalui DatabaseHelper.')

h2('4.3 Tampilan Aplikasi')
para('[ Sertakan screenshot tiap fitur: halaman login, daftar setoran, form tambah, form edit, dialog hapus, pencarian data, dan statistik. Beri keterangan gambar: Gambar 4.1 Tampilan Login, Gambar 4.2 Daftar Setoran, dll. ]', italic=True)

h2('4.4 Pengujian')
para('Berikut adalah hasil pengujian fungsi-fungsi utama aplikasi:')

add_table(
    ['No', 'Skenario Uji', 'Hasil Diharapkan', 'Status'],
    [
        ['1', 'Login dengan credential valid (admin/12345)', 'Berhasil masuk ke MainActivity', 'Berhasil'],
        ['2', 'Login dengan credential kosong', 'Muncul pesan error "Kolom wajib diisi"', 'Berhasil'],
        ['3', 'Login dengan credential salah', 'Muncul pesan error "Kombinasi tidak cocok"', 'Berhasil'],
        ['4', 'Tambah data setoran baru', 'Data masuk dan tampil di daftar', 'Berhasil'],
        ['5', 'Tambah data dengan field kosong', 'Muncul pesan validasi error', 'Berhasil'],
        ['6', 'Ubah data setoran', 'Data ter-update di database', 'Berhasil'],
        ['7', 'Hapus data setoran', 'Data terhapus dari daftar', 'Berhasil'],
        ['8', 'Cari data berdasarkan nama', 'Menampilkan data sesuai kata kunci', 'Berhasil'],
        ['9', 'Urutkan data berdasarkan nama A-Z', 'Data terurut secara ascending', 'Berhasil'],
        ['10', 'Logout dari aplikasi', 'Kembali ke halaman login', 'Berhasil'],
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════
# BAB V
# ═══════════════════════════════════════════
h1('BAB V PENUTUP')

h2('5.1 Kesimpulan')
para('Berdasarkan pengembangan dan pengujian yang telah dilakukan, dapat disimpulkan bahwa aplikasi Bank Sampah telah berhasil dibuat dan dapat berfungsi dengan baik. Aplikasi ini mampu melakukan operasi CRUD pada data setoran sampah, dilengkapi dengan fitur pencarian, pengurutan, autentikasi pengguna, serta tampilan antarmuka yang menggunakan pedoman Material Design.')
para('Konsep-konsep pemrograman mobile yang diterapkan dalam aplikasi ini meliputi pemrograman berorientasi objek, SQLite sebagai media penyimpanan data lokal, RecyclerView untuk tampilan daftar efisien, Activity lifecycle, serta SharedPreferences untuk penyimpanan data sesi. Seluruh fitur utama telah teruji dan berjalan sesuai harapan.')

h2('5.2 Saran')
para('Beberapa saran untuk pengembangan aplikasi ini di masa mendatang adalah:')
numbered('Menambahkan autentikasi berbasis database sehingga pengguna dapat membuat akun sendiri secara mandiri.', restart=True)
numbered('Mengintegrasikan koneksi ke server backend (REST API) agar data dapat diakses dari beberapa perangkat secara bersamaan.')
numbered('Menambahkan fitur laporan dan grafik untuk menampilkan statistik setoran dalam periode waktu tertentu.')
numbered('Menerapkan kalkulasi otomatis saldo berdasarkan jenis dan berat sampah sesuai harga pasar.')
numbered('Menambahkan fitur backup dan restore data untuk keamanan data pengguna.')

doc.add_page_break()


# ═══════════════════════════════════════════
# DAFTAR PUSTAKA
# ═══════════════════════════════════════════
center('DAFTAR PUSTAKA', 14, True)

for ref in [
    'Deitel, P. & Deitel, H. (2017). Java How to Program. Pearson.',
    'Google Developers. (2024). Android App Development Documentation. https://developer.android.com',
    'SQLite Consortium. (2024). SQLite Documentation. https://www.sqlite.org/docs.html',
    'Material Design. (2024). Material Design Guidelines. https://m3.material.io',
    'Android Studio. (2024). Android Studio User Guide. https://developer.android.com/studio',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

doc.add_page_break()


# ═══════════════════════════════════════════
# LAMPIRAN
# ═══════════════════════════════════════════
center('LAMPIRAN', 14, True)
para('a. Link Repository GitHub:', bold=True, indent_first=False)
para('[ https://github.com/username/bank-sampah ]', italic=True, indent_first=False)
para('b. File database: db_sampah_231011400285.db', bold=True, indent_first=False)
para('Terdapat pada direktori data aplikasi di perangkat Android.', indent_first=False)


# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Laporan_Tugas_Akhir_Bank_Sampah.docx')
doc.save(output_path)
print(f'OK: {output_path}')
